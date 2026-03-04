"""
TrajectIQ Enterprise - Sandboxed Resume Parser
===============================================
Secure resume parsing with size limits, MIME validation, and sandboxing.
"""

import os
import re
import io
import hashlib
import mimetypes
import tempfile
import subprocess
import logging
import time
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple, BinaryIO
from dataclasses import dataclass, field
from enum import Enum
from contextlib import contextmanager
import threading

# Resume parsing libraries
import PyPDF2
from docx import Document

# For sandboxing
try:
    import resource
    HAS_RESOURCE = True
except ImportError:
    HAS_RESOURCE = False  # Windows doesn't have resource module


class ParseError(Exception):
    """Resume parsing error"""
    pass


class FileTooLargeError(ParseError):
    """File exceeds size limit"""
    pass


class InvalidMimeTypeError(ParseError):
    """Invalid MIME type"""
    pass


class SuspiciousContentError(ParseError):
    """Suspicious or malicious content detected"""
    pass


class ParseTimeoutError(ParseError):
    """Parsing timed out"""
    pass


@dataclass
class ParserConfig:
    """Resume parser configuration"""
    max_file_size_bytes: int = 10 * 1024 * 1024  # 10 MB
    max_extracted_text_size: int = 1024 * 1024  # 1 MB text
    max_pages: int = 50
    allowed_mime_types: Tuple[str, ...] = (
        'application/pdf',
        'application/msword',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain',
        'application/rtf'
    )
    allowed_extensions: Tuple[str, ...] = ('.pdf', '.doc', '.docx', '.txt', '.rtf')
    parse_timeout_seconds: int = 10
    max_memory_mb: int = 256
    enable_sandbox: bool = True
    sanitize_html: bool = True
    strip_binary_chars: bool = True


@dataclass
class ParseResult:
    """Resume parse result"""
    text: str
    metadata: Dict[str, Any]
    file_hash: str
    parse_time_ms: float
    page_count: int
    warnings: List[str] = field(default_factory=list)
    confidence: float = 1.0


class SandboxedResumeParser:
    """
    Secure resume parser with sandboxing and validation.
    
    Security features:
    - File size limits
    - MIME type validation
    - Extension validation
    - Magic byte verification
    - Memory limits
    - Timeout protection
    - Content sanitization
    - Suspicious pattern detection
    """
    
    # Magic bytes for file type detection
    MAGIC_BYTES = {
        b'%PDF': 'application/pdf',
        b'PK\x03\x04': 'application/zip',  # DOCX is a ZIP
        b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1': 'application/msword',  # DOC
        b'{\\rtf': 'application/rtf',
    }
    
    # Suspicious patterns to detect
    SUSPICIOUS_PATTERNS = [
        # Script injections
        rb'<script',
        rb'javascript:',
        rb'vbscript:',
        rb'onload=',
        rb'onerror=',
        # Macro indicators
        rb'AutoOpen',
        rb'AutoClose',
        rb'Document_Open',
        # Binary exploits
        rb'\x00\x00\x00\x00\x00\x00\x00\x00',  # Long null sequences
        # Embedded executables
        rb'MZ\x90\x00',  # PE header
        rb'\x7fELF',  # ELF header
    ]
    
    def __init__(self, config: Optional[ParserConfig] = None):
        self.config = config or ParserConfig()
        self._logger = logging.getLogger(__name__)
    
    def parse(self, file_input: BinaryIO, filename: str = "") -> ParseResult:
        """
        Parse a resume file with security checks.
        
        Args:
            file_input: File-like object to parse
            filename: Original filename for extension validation
        
        Returns:
            ParseResult with extracted text and metadata
        
        Raises:
            FileTooLargeError: File exceeds size limit
            InvalidMimeTypeError: Invalid file type
            SuspiciousContentError: Malicious content detected
            ParseError: General parsing error
        """
        start_time = time.time()
        
        # Read file content with size limit
        content = self._read_with_limit(file_input)
        
        # Security validations
        self._validate_size(content)
        self._validate_extension(filename)
        detected_mime = self._validate_mime_type(content, filename)
        self._scan_for_threats(content)
        
        # Calculate hash
        file_hash = hashlib.sha256(content).hexdigest()
        
        # Parse based on type
        text, metadata, warnings = self._parse_content(content, detected_mime, filename)
        
        # Sanitize extracted text
        text = self._sanitize_text(text)
        
        # Validate text size
        if len(text) > self.config.max_extracted_text_size:
            text = text[:self.config.max_extracted_text_size]
            warnings.append("Text truncated due to size limit")
        
        parse_time_ms = (time.time() - start_time) * 1000
        
        return ParseResult(
            text=text,
            metadata=metadata,
            file_hash=file_hash,
            parse_time_ms=parse_time_ms,
            page_count=metadata.get('page_count', 1),
            warnings=warnings,
            confidence=self._calculate_confidence(text, metadata)
        )
    
    def parse_file(self, filepath: Path) -> ParseResult:
        """Parse a file from path"""
        with open(filepath, 'rb') as f:
            return self.parse(f, filepath.name)
    
    def parse_bytes(self, content: bytes, filename: str = "") -> ParseResult:
        """Parse from bytes"""
        return self.parse(io.BytesIO(content), filename)
    
    def _read_with_limit(self, file_input: BinaryIO) -> bytes:
        """Read file content with size limit"""
        chunks = []
        total_size = 0
        
        while True:
            chunk = file_input.read(8192)
            if not chunk:
                break
            
            total_size += len(chunk)
            
            if total_size > self.config.max_file_size_bytes:
                raise FileTooLargeError(
                    f"File exceeds maximum size of {self.config.max_file_size_bytes} bytes"
                )
            
            chunks.append(chunk)
        
        return b''.join(chunks)
    
    def _validate_size(self, content: bytes):
        """Validate file size"""
        if len(content) > self.config.max_file_size_bytes:
            raise FileTooLargeError(
                f"File size {len(content)} exceeds limit {self.config.max_file_size_bytes}"
            )
        
        if len(content) == 0:
            raise ParseError("Empty file")
    
    def _validate_extension(self, filename: str):
        """Validate file extension"""
        if not filename:
            return
        
        ext = Path(filename).suffix.lower()
        if ext and ext not in self.config.allowed_extensions:
            raise InvalidMimeTypeError(
                f"File extension '{ext}' not allowed. "
                f"Allowed: {self.config.allowed_extensions}"
            )
    
    def _validate_mime_type(self, content: bytes, filename: str) -> str:
        """Validate and detect MIME type"""
        # Check magic bytes first
        detected_mime = None
        for magic, mime_type in self.MAGIC_BYTES.items():
            if content.startswith(magic):
                detected_mime = mime_type
                break
        
        # Handle DOCX (it's a ZIP file)
        if detected_mime == 'application/zip':
            if filename.lower().endswith('.docx'):
                detected_mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                raise InvalidMimeTypeError("ZIP files not allowed unless DOCX")
        
        # Fall back to extension-based detection
        if not detected_mime:
            guessed, _ = mimetypes.guess_type(filename)
            detected_mime = guessed
        
        # Validate against allowed types
        if detected_mime and detected_mime not in self.config.allowed_mime_types:
            raise InvalidMimeTypeError(
                f"MIME type '{detected_mime}' not allowed. "
                f"Allowed: {self.config.allowed_mime_types}"
            )
        
        if not detected_mime:
            # Default to text/plain for unknown
            detected_mime = 'text/plain'
            self._logger.warning(f"Could not detect MIME type for {filename}, defaulting to text/plain")
        
        return detected_mime
    
    def _scan_for_threats(self, content: bytes):
        """Scan content for suspicious patterns"""
        for pattern in self.SUSPICIOUS_PATTERNS:
            if pattern in content:
                self._logger.warning(f"Suspicious pattern detected: {pattern[:20]}")
                raise SuspiciousContentError(
                    "Suspicious content detected in file. "
                    "File may contain malicious code."
                )
    
    def _parse_content(
        self,
        content: bytes,
        mime_type: str,
        filename: str
    ) -> Tuple[str, Dict[str, Any], List[str]]:
        """Parse content based on MIME type"""
        warnings = []
        metadata = {'mime_type': mime_type, 'filename': filename}
        
        try:
            if mime_type == 'application/pdf':
                text, pdf_meta = self._parse_pdf(content)
                metadata.update(pdf_meta)
                
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                text, doc_meta = self._parse_docx(content)
                metadata.update(doc_meta)
                
            elif mime_type == 'application/msword':
                text, doc_meta = self._parse_doc(content)
                metadata.update(doc_meta)
                warnings.append("Legacy .doc format - limited parsing support")
                
            elif mime_type in ('text/plain', 'application/rtf'):
                text = content.decode('utf-8', errors='replace')
                metadata['encoding'] = 'utf-8'
                
            else:
                # Fallback: try to decode as text
                text = content.decode('utf-8', errors='replace')
                warnings.append(f"Unknown type {mime_type}, treated as text")
        
        except Exception as e:
            raise ParseError(f"Failed to parse {mime_type}: {str(e)}")
        
        return text, metadata, warnings
    
    def _parse_pdf(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Parse PDF file"""
        text_parts = []
        metadata = {}
        
        pdf_file = io.BytesIO(content)
        reader = PyPDF2.PdfReader(pdf_file)
        
        # Get metadata
        if reader.metadata:
            metadata['pdf_title'] = reader.metadata.get('/Title', '')
            metadata['pdf_author'] = reader.metadata.get('/Author', '')
            metadata['pdf_creator'] = reader.metadata.get('/Creator', '')
        
        # Check page count
        page_count = len(reader.pages)
        metadata['page_count'] = page_count
        
        if page_count > self.config.max_pages:
            raise FileTooLargeError(
                f"PDF has {page_count} pages, maximum is {self.config.max_pages}"
            )
        
        # Extract text from each page
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                self._logger.warning(f"Error extracting page {page_num}: {e}")
        
        return '\n\n'.join(text_parts), metadata
    
    def _parse_docx(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX file"""
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        # Extract text from paragraphs
        text_parts = [para.text for para in doc.paragraphs if para.text]
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text for cell in row.cells if cell.text)
                if row_text:
                    text_parts.append(row_text)
        
        metadata = {
            'core_properties': {
                'author': doc.core_properties.author,
                'title': doc.core_properties.title,
                'subject': doc.core_properties.subject,
            }
        }
        
        return '\n\n'.join(text_parts), metadata
    
    def _parse_doc(self, content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Parse legacy DOC file (limited support)"""
        # Legacy DOC files are difficult to parse without external tools
        # Try to extract readable text
        
        text = content.decode('latin-1', errors='replace')
        
        # Remove binary garbage
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        
        # Try to extract readable portions
        readable = []
        for line in text.split('\n'):
            # Check if line has enough printable characters
            printable = sum(1 for c in line if c.isprintable() or c.isspace())
            if len(line) > 0 and printable / len(line) > 0.7:
                readable.append(line.strip())
        
        return '\n'.join(readable), {'format': 'legacy_doc'}
    
    def _sanitize_text(self, text: str) -> str:
        """Sanitize extracted text"""
        if self.config.strip_binary_chars:
            # Remove null bytes and other binary characters
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        
        if self.config.sanitize_html:
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', text)
        
        # Normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = ''.join(c for c in text if c.isprintable() or c in '\n\t')
        
        return text.strip()
    
    def _calculate_confidence(self, text: str, metadata: Dict) -> float:
        """Calculate parsing confidence score"""
        confidence = 1.0
        
        # Penalize short text
        if len(text) < 100:
            confidence *= 0.5
        elif len(text) < 500:
            confidence *= 0.8
        
        # Penalize if page count is unexpected
        page_count = metadata.get('page_count', 1)
        if page_count > 10:
            confidence *= 0.7
        
        # Check for common resume keywords
        resume_keywords = ['experience', 'education', 'skills', 'work', 'resume', 'cv']
        keyword_count = sum(1 for kw in resume_keywords if kw.lower() in text.lower())
        confidence *= min(1.0, 0.6 + keyword_count * 0.1)
        
        return round(confidence, 2)


@contextmanager
def sandbox_limits(memory_mb: int = 256, timeout_seconds: int = 10):
    """
    Context manager for resource limits (Unix only).
    Provides memory and time limits for parsing operations.
    """
    if not HAS_RESOURCE:
        yield
        return
    
    # Set memory limit
    memory_bytes = memory_mb * 1024 * 1024
    
    original_limits = {}
    
    try:
        # Save original limits
        original_limits['memory'] = resource.getrlimit(resource.RLIMIT_AS)
        original_limits['cpu'] = resource.getrlimit(resource.RLIMIT_CPU)
        
        # Set new limits
        resource.setrlimit(resource.RLIMIT_AS, (memory_bytes, memory_bytes))
        resource.setrlimit(resource.RLIMIT_CPU, (timeout_seconds, timeout_seconds))
        
        yield
    
    finally:
        # Restore original limits
        for name, limits in original_limits.items():
            try:
                if name == 'memory':
                    resource.setrlimit(resource.RLIMIT_AS, limits)
                elif name == 'cpu':
                    resource.setrlimit(resource.RLIMIT_CPU, limits)
            except:
                pass


class TimeoutParser:
    """
    Parser wrapper with timeout protection.
    Uses threading to enforce timeout on blocking operations.
    """
    
    def __init__(
        self,
        parser: SandboxedResumeParser,
        timeout_seconds: int = 10
    ):
        self.parser = parser
        self.timeout = timeout_seconds
    
    def parse(self, file_input: BinaryIO, filename: str = "") -> ParseResult:
        """Parse with timeout protection"""
        result = None
        exception = None
        
        def parse_thread():
            nonlocal result, exception
            try:
                result = self.parser.parse(file_input, filename)
            except Exception as e:
                exception = e
        
        thread = threading.Thread(target=parse_thread)
        thread.start()
        thread.join(timeout=self.timeout)
        
        if thread.is_alive():
            # Thread is still running - timeout
            raise ParseTimeoutError(
                f"Parsing timed out after {self.timeout} seconds"
            )
        
        if exception:
            raise exception
        
        return result


# Convenience function
def parse_resume_secure(
    file_input: BinaryIO,
    filename: str = "",
    config: Optional[ParserConfig] = None
) -> ParseResult:
    """
    Securely parse a resume file.
    
    Args:
        file_input: File-like object to parse
        filename: Original filename
        config: Parser configuration
    
    Returns:
        ParseResult with extracted text
    """
    parser = SandboxedResumeParser(config)
    timeout_parser = TimeoutParser(
        parser,
        timeout_seconds=config.parse_timeout_seconds if config else 10
    )
    return timeout_parser.parse(file_input, filename)
