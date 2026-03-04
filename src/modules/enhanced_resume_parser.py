"""
TrajectIQ Enhanced Resume Parser
================================
Advanced resume parsing with OCR, multi-language support, and improved extraction.

Features:
- PDF/DOCX parsing with OCR fallback
- Multi-language detection and support
- Section detection and classification
- Table/chart extraction
- Contact information parsing
- Skills taxonomy mapping
"""

import re
import os
import logging
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from abc import ABC, abstractmethod

# Optional imports for enhanced features
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from langdetect import detect, LangDetectException
    LANGDETECT_AVAILABLE = True
except ImportError:
    LANGDETECT_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


@dataclass
class ParsedResume:
    """Parsed resume data structure"""
    raw_text: str
    contact_info: Dict[str, Any]
    summary: str
    skills: List[str]
    experience: List[Dict[str, Any]]
    education: List[Dict[str, Any]]
    certifications: List[str]
    languages: List[str]
    sections: Dict[str, str]
    metadata: Dict[str, Any] = field(default_factory=dict)
    parsing_errors: List[str] = field(default_factory=list)
    language_detected: str = "en"
    confidence_scores: Dict[str, float] = field(default_factory=dict)


@dataclass
class SectionInfo:
    """Information about a detected section"""
    name: str
    start_index: int
    end_index: int
    content: str
    confidence: float


class ResumeParser:
    """
    Enhanced resume parser with multi-format support.
    """

    # Section detection patterns (multi-language)
    SECTION_PATTERNS = {
        "summary": [
            r"(?i)(professional\s+)?summary",
            r"(?i)objective",
            r"(?i)profile",
            r"(?i)about\s+me",
            r"(?i)career\s+summary",
            r"(?i)resumen\s+profesional",  # Spanish
            r"(?i)profil\s+professionnel",  # French
            r"(?i)profilo\s+professionale",  # Italian
            r"(?i)berufsprofil",  # German
        ],
        "experience": [
            r"(?i)(work\s+)?experience",
            r"(?i)employment\s+history",
            r"(?i)professional\s+experience",
            r"(?i)work\s+history",
            r"(?i)career\s+history",
            r"(?i)experiencia\s+(laboral|profesional)",  # Spanish
            r"(?i)expérience\s+professionnelle",  # French
            r"(?i)esperienza\s+lavorativa",  # Italian
            r"(?i)berufserfahrung",  # German
        ],
        "education": [
            r"(?i)education",
            r"(?i)academic\s+background",
            r"(?i)qualifications",
            r"(?i)academic\s+qualifications",
            r"(?i)educación",  # Spanish
            r"(?i)formation",  # French
            r"(?i)istruzione",  # Italian
            r"(?i)ausbildung",  # German
        ],
        "skills": [
            r"(?i)skills",
            r"(?i)technical\s+skills",
            r"(?i)competencies",
            r"(?i)expertise",
            r"(?i)habilidades",  # Spanish
            r"(?i)compétences",  # French
            r"(?i)competenze",  # Italian
            r"(?i)kenntnisse",  # German
        ],
        "certifications": [
            r"(?i)certifications?",
            r"(?i)licenses?",
            r"(?i)professional\s+development",
            r"(?i)certificaciones",  # Spanish
            r"(?i)certifications",  # French
            r"(?i)certificazioni",  # Italian
            r"(?i)zertifikate",  # German
        ],
        "languages": [
            r"(?i)languages?",
            r"(?i)language\s+skills",
            r"(?i)idiomas",  # Spanish
            r"(?i)langues",  # French
            r"(?i)lingue",  # Italian
            r"(?i)sprachen",  # German
        ],
        "projects": [
            r"(?i)projects?",
            r"(?i)key\s+projects",
            r"(?i)notable\s+projects",
            r"(?i)proyectos",  # Spanish
            r"(?i)projets",  # French
            r"(?i)progetti",  # Italian
            r"(?i)projekte",  # German
        ]
    }

    # Contact info patterns
    CONTACT_PATTERNS = {
        "email": r'[\w\.-]+@[\w\.-]+\.\w+',
        "phone": r'(?:\+\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}',
        "linkedin": r'(?:linkedin\.com/in/|linkedin:)\s*([\w-]+)',
        "github": r'(?:github\.com/|github:)\s*([\w-]+)',
        "website": r'https?://(?:www\.)?[\w\.-]+\.\w+',
    }

    # Skills taxonomy for mapping
    SKILLS_TAXONOMY = {
        # Programming languages
        "programming": [
            "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
            "ruby", "php", "swift", "kotlin", "scala", "r", "matlab", "perl"
        ],
        # Web development
        "web": [
            "react", "angular", "vue", "node.js", "django", "flask", "fastapi",
            "express", "spring", "asp.net", "ruby on rails", "laravel", "symfony"
        ],
        # Databases
        "database": [
            "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
            "oracle", "sql server", "sqlite", "cassandra", "dynamodb", "neo4j"
        ],
        # Cloud & DevOps
        "cloud": [
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
            "jenkins", "ci/cd", "git", "github", "gitlab", "bitbucket"
        ],
        # Data & AI
        "data": [
            "machine learning", "deep learning", "nlp", "computer vision",
            "tensorflow", "pytorch", "pandas", "numpy", "spark", "hadoop",
            "tableau", "power bi", "data science"
        ],
        # Soft skills
        "soft": [
            "leadership", "communication", "teamwork", "problem-solving",
            "project management", "agile", "scrum", "critical thinking"
        ]
    }

    def __init__(self, enable_ocr: bool = True, enable_language_detection: bool = True):
        """
        Initialize parser with feature flags.

        Args:
            enable_ocr: Enable OCR for image-based resumes
            enable_language_detection: Enable automatic language detection
        """
        self.enable_ocr = enable_ocr and OCR_AVAILABLE
        self.enable_language_detection = enable_language_detection and LANGDETECT_AVAILABLE
        self.logger = logging.getLogger(__name__)

    def parse(self, file_path: str) -> ParsedResume:
        """
        Parse a resume file.

        Args:
            file_path: Path to resume file (PDF, DOCX, TXT)

        Returns:
            ParsedResume with extracted information
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        # Extract text based on file type
        text = self._extract_text(path)

        if not text.strip():
            # Try OCR if text extraction failed
            if self.enable_ocr and path.suffix.lower() == '.pdf':
                text = self._ocr_extract(path)

        if not text.strip():
            return ParsedResume(
                raw_text="",
                contact_info={},
                summary="",
                skills=[],
                experience=[],
                education=[],
                certifications=[],
                languages=[],
                sections={},
                parsing_errors=["Could not extract text from file"]
            )

        # Detect language
        language = "en"
        if self.enable_language_detection:
            try:
                language = detect(text)
            except Exception:
                pass

        # Parse sections
        sections = self._parse_sections(text)

        # Extract contact info
        contact_info = self._extract_contact_info(text)

        # Extract skills
        skills = self._extract_skills(text, sections.get("skills", ""))

        # Extract experience
        experience = self._extract_experience(sections.get("experience", ""))

        # Extract education
        education = self._extract_education(sections.get("education", ""))

        # Extract certifications
        certifications = self._extract_certifications(sections.get("certifications", ""))

        # Extract languages
        languages = self._extract_languages(sections.get("languages", ""))

        # Generate confidence scores
        confidence_scores = self._calculate_confidence(
            text, sections, contact_info, skills, experience, education
        )

        return ParsedResume(
            raw_text=text,
            contact_info=contact_info,
            summary=sections.get("summary", ""),
            skills=skills,
            experience=experience,
            education=education,
            certifications=certifications,
            languages=languages,
            sections=sections,
            metadata={
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "parsed_at": datetime.now().isoformat(),
                "parser_version": "3.0.3"
            },
            language_detected=language,
            confidence_scores=confidence_scores
        )

    def _extract_text(self, path: Path) -> str:
        """Extract text from various file formats"""
        suffix = path.suffix.lower()

        if suffix == '.pdf' and PDF_AVAILABLE:
            return self._extract_pdf_text(path)
        elif suffix in ['.docx', '.doc'] and DOCX_AVAILABLE:
            return self._extract_docx_text(path)
        elif suffix == '.txt':
            return path.read_text(encoding='utf-8', errors='ignore')
        else:
            # Fallback: try to read as text
            try:
                return path.read_text(encoding='utf-8', errors='ignore')
            except Exception:
                return ""

    def _extract_pdf_text(self, path: Path) -> str:
        """Extract text from PDF file"""
        text_parts = []

        try:
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            self.logger.warning(f"PDF extraction error: {e}")

        return "\n\n".join(text_parts)

    def _extract_docx_text(self, path: Path) -> str:
        """Extract text from DOCX file"""
        text_parts = []

        try:
            doc = Document(str(path))

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text_parts.append(row_text)
        except Exception as e:
            self.logger.warning(f"DOCX extraction error: {e}")

        return "\n\n".join(text_parts)

    def _ocr_extract(self, path: Path) -> str:
        """Extract text using OCR"""
        if not self.enable_ocr:
            return ""

        text_parts = []

        try:
            images = convert_from_path(str(path))

            for image in images:
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    text_parts.append(page_text)
        except Exception as e:
            self.logger.warning(f"OCR extraction error: {e}")

        return "\n\n".join(text_parts)

    def _parse_sections(self, text: str) -> Dict[str, str]:
        """Parse resume into sections"""
        sections = {}
        lines = text.split('\n')

        current_section = None
        current_content = []

        for line in lines:
            line_stripped = line.strip()

            # Check if line is a section header
            detected_section = self._detect_section_header(line_stripped)

            if detected_section:
                # Save previous section
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content).strip()

                current_section = detected_section
                current_content = []
            elif current_section:
                current_content.append(line)

        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()

        return sections

    def _detect_section_header(self, line: str) -> Optional[str]:
        """Detect if line is a section header"""
        for section, patterns in self.SECTION_PATTERNS.items():
            for pattern in patterns:
                if re.match(pattern + r'\s*$', line):
                    return section
                if re.match(pattern + r'\s*[:|-]?\s*$', line):
                    return section
        return None

    def _extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information"""
        contact = {}

        for field, pattern in self.CONTACT_PATTERNS.items():
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                contact[field] = matches[0] if isinstance(matches[0], str) else matches[0][0]

        # Extract name (usually first non-empty line)
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        if lines:
            # Check if first line looks like a name (2-4 words, mostly letters)
            first_line = lines[0]
            if re.match(r'^[A-Za-z\s\.-]{2,40}$', first_line):
                contact['name'] = first_line

        return contact

    def _extract_skills(self, text: str, skills_section: str) -> List[str]:
        """Extract skills from text"""
        skills = set()

        # Use skills section if available
        search_text = skills_section if skills_section else text

        # Search for skills in taxonomy
        search_lower = search_text.lower()

        for category, skill_list in self.SKILLS_TAXONOMY.items():
            for skill in skill_list:
                if skill.lower() in search_lower:
                    skills.add(skill)

        # Extract comma/semicolon separated lists
        for sep in [',', ';', '•', '|']:
            if sep in search_text:
                parts = search_text.split(sep)
                for part in parts:
                    cleaned = part.strip()
                    if 2 < len(cleaned) < 50:  # Reasonable skill length
                        skills.add(cleaned)

        return sorted(list(skills))

    def _extract_experience(self, text: str) -> List[Dict[str, Any]]:
        """Extract work experience entries"""
        experiences = []

        # Pattern for job entries
        job_pattern = r'(?P<title>[\w\s]+)\s*(?:at|@|-|,)\s*(?P<company>[\w\s&]+)'

        for match in re.finditer(job_pattern, text, re.IGNORECASE):
            exp = {
                "title": match.group('title').strip(),
                "company": match.group('company').strip(),
                "description": "",
                "dates": ""
            }

            # Try to extract dates
            date_match = re.search(
                r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}\s*[-–]\s*(?:Present|Current|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}))',
                text[match.end():match.end()+200]
            )
            if date_match:
                exp["dates"] = date_match.group(1)

            experiences.append(exp)

        return experiences[:10]  # Limit to 10 entries

    def _extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extract education entries"""
        education = []

        # Pattern for degree entries
        degree_pattern = r'(?P<degree>(?:Bachelor|Master|PhD|B\.?S\.?|M\.?S\.?|M\.?B\.?A\.?|Associate)[^,\n]*)'

        for match in re.finditer(degree_pattern, text, re.IGNORECASE):
            edu = {
                "degree": match.group('degree').strip(),
                "institution": "",
                "year": ""
            }

            # Try to extract institution
            context = text[match.end():match.end()+100]
            inst_match = re.search(r'(?:at|from)\s+([\w\s]+(?:University|College|Institute|School))', context)
            if inst_match:
                edu["institution"] = inst_match.group(1).strip()

            # Try to extract year
            year_match = re.search(r'\b(19|20)\d{2}\b', context)
            if year_match:
                edu["year"] = year_match.group(0)

            education.append(edu)

        return education[:5]  # Limit to 5 entries

    def _extract_certifications(self, text: str) -> List[str]:
        """Extract certifications"""
        certifications = []

        # Common certification patterns
        cert_patterns = [
            r'(AWS\s+Certified\s+[\w\s]+)',
            r'(Microsoft\s+Certified[\w\s]+)',
            r'(Google\s+Professional\s+[\w\s]+)',
            r'(Certified\s+[\w\s]+)',
            r'([\w\s]+(?:Certification|Certificate))',
        ]

        for pattern in cert_patterns:
            for match in re.finditer(pattern, text, re.IGNORECASE):
                cert = match.group(1).strip()
                if cert and cert not in certifications:
                    certifications.append(cert)

        return certifications[:10]

    def _extract_languages(self, text: str) -> List[str]:
        """Extract language skills"""
        languages = []

        # Common languages
        known_languages = [
            "english", "spanish", "french", "german", "italian", "portuguese",
            "chinese", "japanese", "korean", "russian", "arabic", "hindi",
            "dutch", "swedish", "polish", "turkish", "vietnamese", "thai"
        ]

        text_lower = text.lower()
        for lang in known_languages:
            if lang in text_lower:
                languages.append(lang.title())

        return languages

    def _calculate_confidence(
        self,
        text: str,
        sections: Dict[str, str],
        contact_info: Dict,
        skills: List,
        experience: List,
        education: List
    ) -> Dict[str, float]:
        """Calculate confidence scores for extracted data"""
        scores = {}

        # Contact info confidence
        contact_fields = ["email", "phone", "name"]
        found_contact = sum(1 for f in contact_fields if f in contact_info)
        scores["contact"] = found_contact / len(contact_fields)

        # Skills confidence
        if len(skills) >= 5:
            scores["skills"] = 0.9
        elif len(skills) >= 3:
            scores["skills"] = 0.7
        else:
            scores["skills"] = 0.4

        # Experience confidence
        if len(experience) >= 2:
            scores["experience"] = 0.8
        elif len(experience) >= 1:
            scores["experience"] = 0.6
        else:
            scores["experience"] = 0.3

        # Education confidence
        if len(education) >= 1:
            scores["education"] = 0.8
        else:
            scores["education"] = 0.4

        # Overall confidence
        scores["overall"] = sum(scores.values()) / len(scores)

        return scores


# Convenience function
def parse_resume(file_path: str, **kwargs) -> ParsedResume:
    """
    Parse a resume file.

    Args:
        file_path: Path to resume file
        **kwargs: Additional parser options

    Returns:
        ParsedResume with extracted information
    """
    parser = ResumeParser(**kwargs)
    return parser.parse(file_path)
