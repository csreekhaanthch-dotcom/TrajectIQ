"""
Docling Document Parser Module
Advanced document parsing for resume intelligence
"""
import asyncio
import json
import hashlib
import time
from typing import Dict, List, Any, Optional
from pathlib import Path
from loguru import logger
from pydantic import BaseModel

# Document parsing
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.pipeline.simple_pipeline import SimplePipeline
    from docling.pipeline.standard_pdf_pipeline import StandardPdfPipeline
    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False
    logger.warning("Docling not available. Install with: pip install docling")


class ParsedResume(BaseModel):
    """Structured resume data from Docling parsing"""
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    
    # Structured sections
    summary: Optional[str] = None
    skills: List[str] = []
    experience: List[Dict[str, Any]] = []
    education: List[Dict[str, Any]] = []
    certifications: List[Dict[str, Any]] = []
    projects: List[Dict[str, Any]] = []
    
    # Raw text
    full_text: str = ""
    
    # Document structure
    sections: Dict[str, str] = {}
    tables: List[Dict[str, Any]] = []
    
    # Metadata
    page_count: int = 0
    word_count: int = 0
    parse_time_ms: int = 0
    confidence: float = 1.0


class DoclingParser:
    """
    Advanced document parser using Docling.
    Handles PDF, DOCX, and other document formats.
    """
    
    def __init__(self):
        self.converter = None
        self._initialized = False
        
    async def initialize(self):
        """Initialize Docling converter"""
        if not DOCLING_AVAILABLE:
            logger.warning("Docling not available, using fallback parser")
            return
            
        if self._initialized:
            return
            
        try:
            # Configure pipeline for PDF processing
            pipeline_options = PdfPipelineOptions()
            pipeline_options.do_ocr = True
            pipeline_options.do_table_structure = True
            
            self.converter = DocumentConverter(
                format_options={
                    InputFormat.PDF: pipeline_options
                }
            )
            self._initialized = True
            logger.info("Docling parser initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize Docling: {e}")
            self._initialized = False
    
    async def parse_resume(
        self, 
        file_path: Optional[str] = None,
        file_bytes: Optional[bytes] = None,
        file_name: Optional[str] = None
    ) -> ParsedResume:
        """
        Parse a resume document.
        
        Args:
            file_path: Path to the resume file
            file_bytes: Raw bytes of the file
            file_name: Original filename for format detection
            
        Returns:
            ParsedResume with structured data
        """
        start_time = time.time()
        
        # Ensure initialized
        await self.initialize()
        
        try:
            if DOCLING_AVAILABLE and self.converter:
                result = await self._parse_with_docling(file_path, file_bytes, file_name)
            else:
                result = await self._parse_with_fallback(file_path, file_bytes, file_name)
                
            result.parse_time_ms = int((time.time() - start_time) * 1000)
            return result
            
        except Exception as e:
            logger.error(f"Error parsing resume: {e}")
            # Return minimal result
            return ParsedResume(
                full_text="",
                parse_time_ms=int((time.time() - start_time) * 1000),
                confidence=0.0
            )
    
    async def _parse_with_docling(
        self,
        file_path: Optional[str],
        file_bytes: Optional[bytes],
        file_name: Optional[str]
    ) -> ParsedResume:
        """Parse using Docling library"""
        
        # Handle file input
        if file_path:
            source = file_path
        elif file_bytes and file_name:
            # Write to temp file
            import tempfile
            suffix = Path(file_name).suffix
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as f:
                f.write(file_bytes)
                source = f.name
        else:
            raise ValueError("Either file_path or file_bytes+file_name required")
        
        # Convert document
        result = self.converter.convert(source)
        
        # Extract text and structure
        doc = result.document
        full_text = doc.export_to_markdown()
        
        # Parse sections
        sections = {}
        current_section = None
        section_content = []
        
        for item in doc.iterate_items():
            if hasattr(item, 'label') and item.label in ['section_header', 'header']:
                if current_section:
                    sections[current_section] = '\n'.join(section_content)
                current_section = item.text.lower().strip()
                section_content = []
            else:
                if current_section:
                    section_content.append(item.text if hasattr(item, 'text') else str(item))
        
        if current_section:
            sections[current_section] = '\n'.join(section_content)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = {
                'headers': [],
                'rows': []
            }
            if hasattr(table, 'data'):
                table_data['headers'] = table.data.get('headers', [])
                table_data['rows'] = table.data.get('rows', [])
            tables.append(table_data)
        
        # Build structured resume
        return ParsedResume(
            summary=sections.get('summary', sections.get('profile', '')),
            skills=self._extract_skills_from_sections(sections),
            experience=self._parse_experience_section(sections.get('experience', '')),
            education=self._parse_education_section(sections.get('education', '')),
            full_text=full_text,
            sections=sections,
            tables=tables,
            page_count=len(doc.pages) if hasattr(doc, 'pages') else 1,
            word_count=len(full_text.split()),
            confidence=0.9
        )
    
    async def _parse_with_fallback(
        self,
        file_path: Optional[str],
        file_bytes: Optional[bytes],
        file_name: Optional[str]
    ) -> ParsedResume:
        """Fallback parser using PyMuPDF and python-docx"""
        
        full_text = ""
        sections = {}
        
        if file_name and file_name.lower().endswith('.pdf'):
            # PDF parsing with PyMuPDF
            import fitz  # PyMuPDF
            
            if file_path:
                doc = fitz.open(file_path)
            elif file_bytes:
                doc = fitz.open(stream=file_bytes, filetype='pdf')
            else:
                raise ValueError("No input provided")
            
            for page in doc:
                full_text += page.get_text()
            doc.close()
            
        elif file_name and file_name.lower().endswith(('.docx', '.doc')):
            # DOCX parsing
            from docx import Document
            
            if file_path:
                doc = Document(file_path)
            elif file_bytes:
                import io
                doc = Document(io.BytesIO(file_bytes))
            else:
                raise ValueError("No input provided")
            
            for para in doc.paragraphs:
                full_text += para.text + '\n'
        else:
            # Try as plain text
            if file_path:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    full_text = f.read()
            elif file_bytes:
                full_text = file_bytes.decode('utf-8', errors='ignore')
        
        # Parse sections from text
        sections = self._detect_sections(full_text)
        
        return ParsedResume(
            summary=sections.get('summary', ''),
            skills=self._extract_skills_from_sections(sections),
            experience=self._parse_experience_section(sections.get('experience', '')),
            education=self._parse_education_section(sections.get('education', '')),
            full_text=full_text,
            sections=sections,
            tables=[],
            page_count=1,
            word_count=len(full_text.split()),
            confidence=0.7
        )
    
    def _detect_sections(self, text: str) -> Dict[str, str]:
        """Detect and extract sections from resume text"""
        sections = {}
        
        # Common section headers
        section_patterns = [
            ('summary', ['summary', 'profile', 'objective', 'about me', 'professional summary']),
            ('experience', ['experience', 'work experience', 'employment', 'professional experience', 'work history']),
            ('education', ['education', 'academic', 'qualifications', 'academic background']),
            ('skills', ['skills', 'technical skills', 'competencies', 'expertise', 'technologies']),
            ('projects', ['projects', 'personal projects', 'key projects']),
            ('certifications', ['certifications', 'certificates', 'licenses']),
            ('languages', ['languages', 'language skills']),
        ]
        
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_lower = line.strip().lower()
            detected = False
            
            for section_name, patterns in section_patterns:
                if any(pattern in line_lower for pattern in patterns):
                    if current_content:
                        sections[current_section] = '\n'.join(current_content)
                    current_section = section_name
                    current_content = []
                    detected = True
                    break
            
            if not detected:
                current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_skills_from_sections(self, sections: Dict[str, str]) -> List[str]:
        """Extract skills from skills section"""
        skills_text = sections.get('skills', '')
        if not skills_text:
            return []
        
        # Common delimiters
        skills = []
        for delimiter in [',', '•', '|', '\n', ';']:
            if delimiter in skills_text:
                skills = [s.strip() for s in skills_text.split(delimiter) if s.strip()]
                break
        
        if not skills:
            skills = [skills_text.strip()]
        
        return skills[:50]  # Limit to 50 skills
    
    def _parse_experience_section(self, text: str) -> List[Dict[str, Any]]:
        """Parse experience entries"""
        if not text:
            return []
        
        experiences = []
        lines = text.strip().split('\n')
        current_exp = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Detect company/title patterns
            if any(word in line.lower() for word in ['inc', 'llc', 'corp', 'ltd', 'company']):
                if current_exp:
                    experiences.append(current_exp)
                current_exp = {'company': line, 'title': '', 'duration': '', 'description': ''}
            elif any(word in line.lower() for word in ['engineer', 'developer', 'manager', 'analyst', 'designer', 'lead']):
                if 'title' not in current_exp or not current_exp.get('title'):
                    current_exp['title'] = line
            elif any(char.isdigit() for char in line) and any(word in line.lower() for word in ['jan', 'feb', 'mar', 'apr', 'may', 'jun', 'jul', 'aug', 'sep', 'oct', 'nov', 'dec', 'present', 'current', '20']):
                current_exp['duration'] = line
            else:
                if current_exp.get('description'):
                    current_exp['description'] += ' ' + line
                else:
                    current_exp['description'] = line
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences[:10]
    
    def _parse_education_section(self, text: str) -> List[Dict[str, Any]]:
        """Parse education entries"""
        if not text:
            return []
        
        education = []
        lines = text.strip().split('\n')
        current_edu = {}
        
        degree_keywords = ['bachelor', 'master', 'phd', 'doctorate', 'mba', 'bs', 'ba', 'ms', 'ma', 'associate', 'diploma']
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if any(keyword in line.lower() for keyword in degree_keywords):
                if current_edu:
                    education.append(current_edu)
                current_edu = {'degree': line, 'institution': '', 'year': ''}
            elif 'university' in line.lower() or 'college' in line.lower() or 'institute' in line.lower():
                if not current_edu.get('institution'):
                    current_edu['institution'] = line
            elif any(char.isdigit() for char in line) and len(line) <= 10:
                current_edu['year'] = line
        
        if current_edu:
            education.append(current_edu)
        
        return education[:5]


# Singleton instance
_parser_instance: Optional[DoclingParser] = None


async def get_parser() -> DoclingParser:
    """Get or create DoclingParser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DoclingParser()
        await _parser_instance.initialize()
    return _parser_instance


async def parse_resume(
    file_path: Optional[str] = None,
    file_bytes: Optional[bytes] = None,
    file_name: Optional[str] = None
) -> ParsedResume:
    """Convenience function to parse a resume"""
    parser = await get_parser()
    return await parser.parse_resume(file_path, file_bytes, file_name)
