"""
TrajectIQ Resume Parser Module
==============================
Parses PDF and DOCX resumes into structured JSON format.
Uses Ollama for intelligent section extraction and entity recognition.
"""

import json
import re
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple
from pathlib import Path

from core.base_module import BaseModule, ModuleRegistry
from core.config import config


@ModuleRegistry.register
class ResumeParser(BaseModule):
    """
    Resume parsing module for PDF and DOCX files.
    Extracts structured information using NLP and heuristics.
    """
    
    module_name = "resume_parser"
    version = "1.0.0"
    
    # Section headers to identify
    SECTION_HEADERS = {
        "experience": ["experience", "work experience", "employment", "work history", "professional experience"],
        "education": ["education", "academic background", "academic history", "qualifications"],
        "skills": ["skills", "technical skills", "competencies", "expertise", "technologies"],
        "projects": ["projects", "personal projects", "key projects", "selected projects"],
        "certifications": ["certifications", "certificates", "professional development"],
        "summary": ["summary", "profile", "objective", "about", "professional summary"],
        "achievements": ["achievements", "accomplishments", "key achievements", "awards"],
        "languages": ["languages", "language skills"],
        "interests": ["interests", "hobbies", "personal interests"]
    }
    
    # Common date patterns
    DATE_PATTERNS = [
        r'(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4})',  # Jan 2020
        r'(\b\d{1,2}/\d{4})',  # 01/2020
        r'(\b\d{4})',  # 2020
        r'((?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})',  # January 2020
    ]
    
    def validate_input(self, input_data: Dict[str, Any]) -> bool:
        """Validate input against schema"""
        if "source_type" not in input_data:
            raise ValueError("Missing required field: source_type")
        
        if "content" not in input_data:
            raise ValueError("Missing required field: content")
        
        valid_source_types = ["file_path", "raw_text", "base64"]
        if input_data["source_type"] not in valid_source_types:
            raise ValueError(f"Invalid source_type. Must be one of: {valid_source_types}")
        
        return True
    
    def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """Process resume and extract structured information"""
        
        # Step 1: Extract raw text
        raw_text = self._extract_text(input_data)
        
        # Step 2: Identify sections
        sections = self._identify_sections(raw_text)
        
        # Step 3: Extract candidate info
        candidate_info = self._extract_candidate_info(sections, raw_text)
        
        # Step 4: Parse education
        education = self._parse_education(sections.get("education", ""))
        
        # Step 5: Parse experience
        experience = self._parse_experience(sections.get("experience", ""))
        
        # Step 6: Extract skills
        skills = self._extract_skills(sections.get("skills", ""), experience)
        
        # Step 7: Parse certifications
        certifications = self._parse_certifications(sections.get("certifications", ""))
        
        # Step 8: Parse projects
        projects = self._parse_projects(sections.get("projects", ""))
        
        # Step 9: Calculate confidence score
        confidence_score = self._calculate_confidence(
            candidate_info, education, experience, skills
        )
        
        # Build output
        output = {
            "parse_id": self.generate_id("PARSE"),
            "timestamp": datetime.utcnow().isoformat() + "Z",
            "status": "success" if confidence_score > 0.5 else "partial",
            "confidence_score": confidence_score,
            "candidate_info": candidate_info,
            "education": education,
            "experience": experience,
            "skills": skills,
            "certifications": certifications,
            "projects": projects,
            "raw_sections": sections,
            "parsing_metadata": {
                "parser_version": self.version,
                "processing_time_ms": 0,  # Will be filled by base class
                "pages_processed": 1,
                "sections_found": list(sections.keys()),
                "warnings": []
            }
        }
        
        # Use LLM for enhancement if available
        if self.ollama_client:
            output = self._llm_enhance(output)
        
        return output
    
    def _extract_text(self, input_data: Dict[str, Any]) -> str:
        """Extract text from various input formats"""
        source_type = input_data["source_type"]
        content = input_data["content"]
        
        if source_type == "raw_text":
            return content
        
        elif source_type == "file_path":
            file_path = Path(content)
            extension = input_data.get("file_extension", file_path.suffix.lstrip("."))
            
            if extension == "pdf":
                return self._extract_from_pdf(file_path)
            elif extension in ["docx", "doc"]:
                return self._extract_from_docx(file_path)
            elif extension == "txt":
                return file_path.read_text(encoding="utf-8", errors="ignore")
            else:
                raise ValueError(f"Unsupported file extension: {extension}")
        
        elif source_type == "base64":
            import base64
            # Decode and process as appropriate format
            decoded = base64.b64decode(content)
            # Try to detect format and extract
            return decoded.decode("utf-8", errors="ignore")
        
        return ""
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            import pdfplumber
            
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            self.logger.warning("pdfplumber not installed, falling back to basic extraction")
            # Fallback to basic text extraction
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text_parts = [page.extract_text() for page in reader.pages]
                return "\n\n".join(text_parts)
            except ImportError:
                self.logger.error("No PDF library available")
                return ""
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            self.logger.error("python-docx not installed")
            return ""
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identify and extract resume sections"""
        sections = {}
        lines = text.split("\n")
        
        current_section = "header"
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            found_section = None
            for section_name, headers in self.SECTION_HEADERS.items():
                if any(header in line_lower for header in headers):
                    # Save previous section
                    if current_content:
                        sections[current_section] = "\n".join(current_content).strip()
                    
                    found_section = section_name
                    current_section = section_name
                    current_content = []
                    break
            
            if not found_section:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content).strip()
        
        return sections
    
    def _extract_candidate_info(
        self,
        sections: Dict[str, str],
        raw_text: str
    ) -> Dict[str, Any]:
        """Extract candidate contact and basic info"""
        info = {
            "full_name": "",
            "email": "",
            "phone": "",
            "location": {},
            "linkedin_url": "",
            "github_url": "",
            "portfolio_url": "",
            "summary": sections.get("summary", "")
        }
        
        # Extract email
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        email_match = re.search(email_pattern, raw_text)
        if email_match:
            info["email"] = email_match.group()
        
        # Extract phone
        phone_pattern = r'(?:\+?1[-.\s]?)?(?:\(?[0-9]{3}\)?[-.\s]?)?[0-9]{3}[-.\s]?[0-9]{4}'
        phone_match = re.search(phone_pattern, raw_text)
        if phone_match:
            info["phone"] = phone_match.group()
        
        # Extract LinkedIn
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, raw_text, re.IGNORECASE)
        if linkedin_match:
            info["linkedin_url"] = linkedin_match.group()
        
        # Extract GitHub
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[\w-]+'
        github_match = re.search(github_pattern, raw_text, re.IGNORECASE)
        if github_match:
            info["github_url"] = github_match.group()
        
        # Extract name (usually first non-empty line or before email)
        header = sections.get("header", "")
        if header:
            # First line is typically the name
            first_lines = [l.strip() for l in header.split("\n") if l.strip()]
            if first_lines:
                # Filter out lines that look like contact info
                for line in first_lines[:3]:
                    if not re.search(email_pattern, line) and not re.search(phone_pattern, line):
                        if not any(url in line.lower() for url in ["linkedin", "github", "http"]):
                            info["full_name"] = line.title()
                            break
        
        return info
    
    def _parse_education(self, education_text: str) -> List[Dict[str, Any]]:
        """Parse education section"""
        if not education_text:
            return []
        
        education = []
        
        # Split by common delimiters
        entries = re.split(r'\n\s*\n|\n(?=[A-Z][^a-z]*(?:University|College|Institute|School))', education_text)
        
        for entry in entries:
            if not entry.strip():
                continue
            
            edu = {
                "institution": "",
                "degree": "",
                "field_of_study": "",
                "start_date": "",
                "end_date": "",
                "gpa": None,
                "honors": [],
                "is_complete": True
            }
            
            # Extract institution
            inst_match = re.search(r'([A-Z][A-Za-z\s]+(?:University|College|Institute|School))', entry)
            if inst_match:
                edu["institution"] = inst_match.group(1).strip()
            
            # Extract degree
            degree_patterns = [
                r'(Bachelor[\'\'\s]*(?:of\s+)?[A-Za-z\s]+)',
                r'(Master[\'\'\s]*(?:of\s+)?[A-Za-z\s]+)',
                r'(PhD|Ph\.?D\.?|Doctorate)',
                r'(MBA)',
                r'(BS|BA|MS|MA|MSc|BSc)\s*(?:in\s+)?([A-Za-z\s]+)?'
            ]
            
            for pattern in degree_patterns:
                match = re.search(pattern, entry, re.IGNORECASE)
                if match:
                    edu["degree"] = match.group(0).strip()
                    break
            
            # Extract dates
            dates = self._extract_dates(entry)
            if dates:
                edu["start_date"] = dates[0]
                if len(dates) > 1:
                    edu["end_date"] = dates[1]
            
            # Extract GPA
            gpa_match = re.search(r'GPA[:\s]*([0-3]\.?[0-9]?)', entry, re.IGNORECASE)
            if gpa_match:
                try:
                    edu["gpa"] = float(gpa_match.group(1))
                except ValueError:
                    pass
            
            # Only add if we found meaningful information
            if edu["institution"] or edu["degree"]:
                education.append(edu)
        
        return education
    
    def _parse_experience(self, experience_text: str) -> List[Dict[str, Any]]:
        """Parse experience section"""
        if not experience_text:
            return []
        
        experience = []
        
        # Split into job entries
        # Look for patterns like "Company Name | Title" or "Title at Company"
        entries = re.split(r'\n\s*\n|\n(?=[A-Z][^a-z]{0,10}(?:Inc|Corp|LLC|Ltd|Company))', experience_text)
        
        for entry in entries:
            if not entry.strip() or len(entry.strip()) < 20:
                continue
            
            job = {
                "company": "",
                "title": "",
                "location": "",
                "start_date": "",
                "end_date": "",
                "is_current": False,
                "employment_type": "full-time",
                "description": "",
                "achievements": [],
                "technologies_used": []
            }
            
            lines = entry.strip().split("\n")
            
            # First line often contains company and title
            if lines:
                first_line = lines[0]
                
                # Try different patterns
                # Pattern: Company | Title
                if "|" in first_line:
                    parts = first_line.split("|")
                    job["company"] = parts[0].strip()
                    if len(parts) > 1:
                        job["title"] = parts[1].strip()
                
                # Pattern: Title at Company
                elif " at " in first_line.lower():
                    parts = re.split(r'\s+at\s+', first_line, flags=re.IGNORECASE)
                    job["title"] = parts[0].strip()
                    if len(parts) > 1:
                        job["company"] = parts[1].strip()
                
                # Pattern: Company - Title
                elif " - " in first_line or " – " in first_line:
                    parts = re.split(r'\s*[-–]\s*', first_line)
                    job["company"] = parts[0].strip()
                    if len(parts) > 1:
                        job["title"] = parts[1].strip()
                else:
                    # Assume first line is title or company
                    job["title"] = first_line.strip()
            
            # Extract dates
            dates = self._extract_dates(entry)
            if dates:
                job["start_date"] = dates[0]
                if len(dates) > 1:
                    job["end_date"] = dates[1]
            
            # Check if current
            if not job["end_date"] or "present" in job["end_date"].lower():
                job["is_current"] = True
            
            # Extract description and achievements
            desc_lines = []
            achievements = []
            
            for line in lines[1:]:
                line = line.strip()
                if not line:
                    continue
                
                # Achievement bullets
                if line.startswith(("•", "-", "*", "►", "→")):
                    achievement_text = line.lstrip("•-*►→ ").strip()
                    
                    # Extract metrics from achievement
                    metrics = self._extract_metrics(achievement_text)
                    
                    achievements.append({
                        "text": achievement_text,
                        "metrics": metrics
                    })
                else:
                    desc_lines.append(line)
            
            job["description"] = " ".join(desc_lines)
            job["achievements"] = achievements
            
            # Extract technologies mentioned
            tech_pattern = r'\b(Python|Java|JavaScript|TypeScript|React|Angular|Vue|Node\.?js|SQL|PostgreSQL|MySQL|MongoDB|Redis|AWS|Azure|GCP|Docker|Kubernetes|Jenkins|Git|Linux|TensorFlow|PyTorch|scikit-learn|Spark|Hadoop|Kafka|GraphQL|REST|API|HTML|CSS|Sass|Tailwind|Bootstrap|Swift|Kotlin|Rust|Go|C\+\+|C#|\.NET|Ruby|Rails|Django|Flask|FastAPI|Spring|Express|NestJS)\b'
            tech_matches = re.findall(tech_pattern, entry, re.IGNORECASE)
            job["technologies_used"] = list(set(t.strip() for t in tech_matches))
            
            # Only add if we found meaningful information
            if job["company"] or job["title"]:
                experience.append(job)
        
        return experience
    
    def _extract_skills(
        self,
        skills_text: str,
        experience: List[Dict]
    ) -> Dict[str, Any]:
        """Extract and categorize skills"""
        skills = {
            "technical": [],
            "soft_skills": [],
            "languages": []
        }
        
        # Technical skill categories
        tech_categories = {
            "Programming": ["python", "java", "javascript", "typescript", "c++", "c#", "go", "rust", "ruby", "php", "swift", "kotlin", "scala", "r", "matlab"],
            "Frontend": ["react", "angular", "vue", "html", "css", "javascript", "typescript", "svelte", "next.js", "tailwind", "bootstrap"],
            "Backend": ["node.js", "django", "flask", "fastapi", "spring", "express", "rails", "asp.net", "graphql", "rest api"],
            "Database": ["sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "cassandra", "dynamodb", "oracle"],
            "Cloud": ["aws", "azure", "gcp", "heroku", "digitalocean", "cloudflare"],
            "DevOps": ["docker", "kubernetes", "jenkins", "git", "ci/cd", "terraform", "ansible", "prometheus", "grafana"],
            "Data Science": ["machine learning", "deep learning", "tensorflow", "pytorch", "pandas", "numpy", "scikit-learn", "spark", "hadoop"],
            "Infrastructure": ["linux", "nginx", "apache", "kafka", "rabbitmq", "microservices", "api gateway"]
        }
        
        # Soft skills
        soft_skill_list = [
            "leadership", "communication", "teamwork", "problem-solving", 
            "analytical", "creative", "adaptability", "time management",
            "project management", "collaboration", "mentoring", "presentation"
        ]
        
        # Known languages
        language_list = [
            "english", "spanish", "french", "german", "chinese", "mandarin",
            "japanese", "korean", "portuguese", "russian", "arabic", "hindi"
        ]
        
        # Combine skills text with experience technologies
        all_text = skills_text.lower()
        for exp in experience:
            all_text += " " + " ".join(exp.get("technologies_used", [])).lower()
        
        # Extract technical skills
        found_skills = set()
        
        for category, skill_list in tech_categories.items():
            for skill in skill_list:
                if skill in all_text:
                    found_skills.add((skill.title(), category))
        
        # Add found technical skills
        for skill_name, category in found_skills:
            # Try to determine proficiency from context
            proficiency = self._determine_proficiency(skill_name.lower(), all_text)
            
            skills["technical"].append({
                "name": skill_name,
                "category": category,
                "years_experience": None,  # Will be filled by skill module
                "proficiency": proficiency
            })
        
        # Extract soft skills
        for skill in soft_skill_list:
            if skill in all_text:
                skills["soft_skills"].append(skill.title())
        
        # Extract languages
        lang_proficiency_map = {
            "native": "Native",
            "fluent": "Fluent",
            "proficient": "Proficient",
            "intermediate": "Intermediate",
            "basic": "Basic",
            "conversational": "Conversational"
        }
        
        for lang in language_list:
            if lang in all_text:
                # Try to find proficiency
                proficiency = "Proficient"  # Default
                for prof_key, prof_val in lang_proficiency_map.items():
                    pattern = rf'{lang}\s*[-:]?\s*{prof_key}'
                    if re.search(pattern, all_text, re.IGNORECASE):
                        proficiency = prof_val
                        break
                
                skills["languages"].append({
                    "language": lang.title(),
                    "proficiency": proficiency
                })
        
        return skills
    
    def _parse_certifications(self, cert_text: str) -> List[Dict[str, Any]]:
        """Parse certifications section"""
        if not cert_text:
            return []
        
        certifications = []
        
        # Common certification patterns
        cert_patterns = [
            r'((?:AWS|Azure|GCP|Google Cloud|Microsoft|Oracle|Cisco|CompTIA|PMP|Scrum|Kubernetes|Docker|Salesforce)[^,\n]+)',
            r'([A-Z]{2,}[\s\-]?(?:Certified|Certification|Professional|Specialist|Expert)[^,\n]+)'
        ]
        
        for pattern in cert_patterns:
            matches = re.finditer(pattern, cert_text, re.IGNORECASE)
            for match in matches:
                cert = {
                    "name": match.group(1).strip(),
                    "issuer": "",
                    "issue_date": "",
                    "expiry_date": "",
                    "credential_id": "",
                    "credential_url": ""
                }
                
                # Try to extract issuer
                if "AWS" in cert["name"]:
                    cert["issuer"] = "Amazon Web Services"
                elif "Azure" in cert["name"] or "Microsoft" in cert["name"]:
                    cert["issuer"] = "Microsoft"
                elif "GCP" in cert["name"] or "Google" in cert["name"]:
                    cert["issuer"] = "Google"
                
                certifications.append(cert)
        
        return certifications
    
    def _parse_projects(self, projects_text: str) -> List[Dict[str, Any]]:
        """Parse projects section"""
        if not projects_text:
            return []
        
        projects = []
        
        # Split into project entries
        entries = re.split(r'\n\s*\n|\n(?=[A-Z][^\n]{5,50}\n)', projects_text)
        
        for entry in entries:
            if not entry.strip() or len(entry.strip()) < 20:
                continue
            
            project = {
                "name": "",
                "description": "",
                "technologies": [],
                "url": "",
                "start_date": "",
                "end_date": ""
            }
            
            lines = entry.strip().split("\n")
            
            if lines:
                project["name"] = lines[0].strip()
                
                # Extract description from remaining lines
                desc_lines = []
                for line in lines[1:]:
                    line = line.strip()
                    if line and not line.startswith(("http", "www")):
                        desc_lines.append(line)
                
                project["description"] = " ".join(desc_lines)
                
                # Extract URL
                url_match = re.search(r'(https?://[^\s]+)', entry)
                if url_match:
                    project["url"] = url_match.group(1)
            
            if project["name"]:
                projects.append(project)
        
        return projects
    
    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text"""
        dates = []
        
        for pattern in self.DATE_PATTERNS:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)
        
        # Normalize dates
        normalized = []
        for date in dates:
            date = date.strip()
            # Add to normalized if it looks like a valid date
            if len(date) >= 4:
                normalized.append(date)
        
        return normalized[:2]  # Return at most 2 dates (start, end)
    
    def _extract_metrics(self, text: str) -> List[Dict[str, str]]:
        """Extract quantifiable metrics from achievement text"""
        metrics = []
        
        # Pattern for percentages
        percent_pattern = r'(\d+(?:\.\d+)?)\s*%'
        for match in re.finditer(percent_pattern, text):
            metrics.append({
                "value": f"{match.group(1)}%",
                "unit": "percentage",
                "context": self._extract_context(text, match.start(), match.end())
            })
        
        # Pattern for multipliers (e.g., 10x, 2x)
        multiplier_pattern = r'(\d+(?:\.\d+)?)\s*x\b'
        for match in re.finditer(multiplier_pattern, text, re.IGNORECASE):
            metrics.append({
                "value": f"{match.group(1)}x",
                "unit": "multiplier",
                "context": self._extract_context(text, match.start(), match.end())
            })
        
        # Pattern for absolute numbers with context (users, customers, etc.)
        number_pattern = r'(\d+(?:,\d{3})*(?:\.\d+)?)\s*(?:million|thousand|k|m|users?|customers?|clients?)'
        for match in re.finditer(number_pattern, text, re.IGNORECASE):
            metrics.append({
                "value": match.group(1),
                "unit": "count",
                "context": match.group(0)
            })
        
        return metrics
    
    def _extract_context(self, text: str, start: int, end: int, window: int = 30) -> str:
        """Extract surrounding context for a metric"""
        context_start = max(0, start - window)
        context_end = min(len(text), end + window)
        return text[context_start:context_end].strip()
    
    def _determine_proficiency(self, skill: str, text: str) -> str:
        """Determine skill proficiency from context"""
        # Look for proficiency indicators
        if re.search(rf'{skill}[\s\-]*(expert|guru|master|senior)', text, re.IGNORECASE):
            return "expert"
        elif re.search(rf'{skill}[\s\-]*(advanced|senior|experienced)', text, re.IGNORECASE):
            return "advanced"
        elif re.search(rf'{skill}[\s\-]*(intermediate|mid)', text, re.IGNORECASE):
            return "intermediate"
        elif re.search(rf'{skill}[\s\-]*(beginner|junior|learning|basic)', text, re.IGNORECASE):
            return "beginner"
        
        # Default based on years mentioned
        years_pattern = rf'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of\s+)?(?:experience\s+(?:with|in|using)\s+)?{skill}'
        match = re.search(years_pattern, text, re.IGNORECASE)
        if match:
            years = int(match.group(1))
            if years >= 7:
                return "expert"
            elif years >= 4:
                return "advanced"
            elif years >= 2:
                return "intermediate"
            else:
                return "beginner"
        
        return "intermediate"  # Default
    
    def _calculate_confidence(
        self,
        candidate_info: Dict,
        education: List,
        experience: List,
        skills: Dict
    ) -> float:
        """Calculate overall parsing confidence score"""
        score = 0.0
        
        # Name extraction (20%)
        if candidate_info.get("full_name"):
            score += 0.15
        if candidate_info.get("email"):
            score += 0.10
        
        # Education (20%)
        if education:
            score += min(0.20, len(education) * 0.10)
        
        # Experience (35%)
        if experience:
            score += min(0.35, len(experience) * 0.12)
        
        # Skills (15%)
        technical_skills = skills.get("technical", [])
        if technical_skills:
            score += min(0.15, len(technical_skills) * 0.02)
        
        return min(1.0, score)
    
    def _llm_enhance(self, output: Dict[str, Any]) -> Dict[str, Any]:
        """Use LLM to enhance parsing results"""
        if not self.ollama_client:
            return output
        
        try:
            # Prepare context for LLM
            raw_sections = output.get("raw_sections", {})
            
            prompt = f"""Analyze this resume data and enhance the structured extraction. Return ONLY valid JSON.

Current extraction:
{json.dumps(output.get("candidate_info", {}), indent=2)}

Raw resume sections:
- Header: {raw_sections.get("header", "")[:500]}
- Experience: {raw_sections.get("experience", "")[:1000]}
- Education: {raw_sections.get("education", "")[:500]}

Please verify and correct:
1. Candidate name (should be properly formatted, not email or phone)
2. Job titles (should be professional titles, not company names)
3. Company names (should be organization names, not titles)
4. Any missing information that can be extracted

Return a JSON object with corrections:
{{
    "candidate_info_corrections": {{
        "full_name": "corrected name if needed",
        "additional_notes": "any notes"
    }},
    "experience_corrections": [
        {{
            "index": 0,
            "title": "corrected title",
            "company": "corrected company"
        }}
    ]
}}"""
            
            response = self.llm_call(prompt)
            corrections = self.ensure_json_output(response)
            
            # Apply corrections
            if corrections.get("candidate_info_corrections"):
                for key, value in corrections["candidate_info_corrections"].items():
                    if key != "additional_notes" and value:
                        output["candidate_info"][key] = value
            
            if corrections.get("experience_corrections"):
                for correction in corrections["experience_corrections"]:
                    idx = correction.get("index", 0)
                    if idx < len(output.get("experience", [])):
                        if correction.get("title"):
                            output["experience"][idx]["title"] = correction["title"]
                        if correction.get("company"):
                            output["experience"][idx]["company"] = correction["company"]
            
        except Exception as e:
            self.logger.warning(f"LLM enhancement failed: {str(e)}")
        
        return output
